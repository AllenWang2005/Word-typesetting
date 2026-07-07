#!/usr/bin/env python3
"""Build a deliberately NON-COMPLIANT sample DOCX to demonstrate the auditor.

Requires python-docx (`pip install python-docx`).

Usage:
    python examples/make_sample.py [output.docx]
    python scripts/audit_docx_format.py output.docx

The generated document intentionally violates several rules so the audit script
has something to report:
  * ASCII punctuation inside Chinese prose
  * a centered level-1 heading
  * a centered / indented abstract paragraph
  * table text that is not small-four (12 pt)
  * full-width and field-less citation brackets
  * a bare superscript citation number with no brackets
  * plain-text formula / visible LaTeX instead of OMML
  * non-black body font color
  * a heading in Songti that ends with punctuation
  * a table caption placed below its table
  * a shaded (non-white) table header cell
"""

from __future__ import annotations

import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def _set_border(parent, tag: str, val: str, sz: str = "") -> None:
    element = OxmlElement(f"w:{tag}")
    element.set(qn("w:val"), val)
    if sz:
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
    parent.append(element)


def make_three_line(table) -> None:
    """Give a python-docx table explicit three-line borders and a white background."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    _set_border(borders, "top", "single", "12")
    _set_border(borders, "bottom", "single", "12")
    for tag in ("left", "right", "insideH", "insideV"):
        _set_border(borders, tag, "none")
    tbl_pr.append(borders)
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        cell_borders = OxmlElement("w:tcBorders")
        _set_border(cell_borders, "bottom", "single", "6")
        tc_pr.append(cell_borders)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "auto")
        tc_pr.append(shd)


def build(path: str) -> None:
    doc = Document()

    # Centered level-1 heading (should be left aligned).
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("一、绪论")

    # Abstract that is centered and indented (should be left aligned, indent 0).
    abs = doc.add_paragraph()
    abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs.paragraph_format.first_line_indent = Pt(24)
    abs.add_run("摘要 本文研究梯级水库的优化调度方法与防洪效益评估。")

    # Body prose with ASCII punctuation in Chinese text, plus a non-black run.
    p = doc.add_paragraph()
    p.add_run("本设计针对某流域的水利计算问题展开分析,采用多方案比较的方法进行评价.")
    red = p.add_run("（本句为红色正文，属于不应保留的彩色字体）")
    red.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Plain-text formula and visible LaTeX (should be native OMML).
    doc.add_paragraph("式中 N_p 为水泵台数，反映系统的装机规模。")
    doc.add_paragraph("两者比值可用 \\frac{a}{b} 表示，需要渲染为公式对象。")

    # Citations with full-width brackets and a field-less ASCII bracket.
    doc.add_paragraph("已有研究表明该方法有效［1］，并在工程中得到验证[2]。")

    # A bare superscript citation number that dropped its brackets (should be [3]).
    bare = doc.add_paragraph()
    bare.add_run("该结论亦见于相关文献")
    sup = bare.add_run("3")
    sup.font.superscript = True
    bare.add_run("。")

    # A heading in Songti (should be Heiti) that ends with punctuation (should not).
    h2 = doc.add_paragraph(style="Heading 1")
    run = h2.add_run("二、研究方法。")
    run.font.name = "宋体"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    # A table whose text is not small-four (12 pt), with a shaded header cell
    # (a three-line table must be all white).
    table = doc.add_table(rows=1, cols=2)
    for idx, label in enumerate(("方案", "流量")):
        cell = table.cell(0, idx)
        run = cell.paragraphs[0].add_run(label)
        run.font.size = Pt(14)  # -> w:sz=28, not the required 24
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "D9E2F3")  # light blue header shading
    table.cell(0, 0)._tc.get_or_add_tcPr().append(shd)

    # A table caption placed BELOW the table (table captions belong above).
    doc.add_paragraph("表 1-1 方案比较结果")

    doc.save(path)
    print(f"wrote {path}")


def build_compliant(path: str) -> None:
    """Build a COMPLIANT sample that the auditor reports as PASS."""
    doc = Document()

    # Cover title: centered Heiti (per spec; the body-font check must not flag it).
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("某某水库水利计算课程设计")
    tr.bold = True
    tr.font.name = "黑体"
    tr._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")

    # Abstract: left aligned, no indent.
    doc.add_paragraph("摘要：本文研究梯级水库的优化调度与防洪效益评估，比较多个方案并给出推荐结论。")

    # Heading 1: real heading style, Heiti, no trailing punctuation, spacing before (no gap warning).
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(12)
    hr = heading.add_run("一、绪论")
    hr.font.name = "黑体"
    hr._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")

    # Body: Chinese prose with Chinese punctuation only.
    doc.add_paragraph("本设计针对某流域的水利计算问题进行分析，比较多个方案的发电与防洪效益，并给出推荐方案。")

    # Table caption ABOVE the table, centered.
    cap = doc.add_paragraph("表 1-1 主要计算参数")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # A proper white three-line table with 五号 (10.5 pt) text.
    table = doc.add_table(rows=2, cols=2)
    for row in table.rows:
        for cell in row.cells:
            run = cell.paragraphs[0].add_run("参数")
            run.font.size = Pt(10.5)  # 五号 -> w:sz=21
    make_three_line(table)

    doc.save(path)
    print(f"wrote {path}")


if __name__ == "__main__":
    cli_args = sys.argv[1:]
    compliant = "--compliant" in cli_args
    paths = [arg for arg in cli_args if not arg.startswith("--")]
    out = paths[0] if paths else ("sample-compliant.docx" if compliant else "sample.docx")
    (build_compliant if compliant else build)(out)
